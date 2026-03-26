"""
Tool: cover_letter.py

Purpose:
    Take a plain-text cover letter written by Claude, format it as a .docx,
    upload it to Google Drive, and return a shareable URL.

    The actual letter content is written by the agent (Claude), not here.
    This tool handles formatting and Drive upload only.

Usage:
    python tools/cover_letter.py \
        --company "Acme Corp" \
        --job_title "Analytics Director" \
        --content_file ".tmp/cover_Acme_AnalyticsDirector.txt"

Parameters:
    --company       Company name (used in filename and letter header)
    --job_title     Job title (used in filename)
    --content_file  Path to plain-text file containing the cover letter body

Content file format:
    Plain text. Paragraphs separated by blank lines.
    Do not include a salutation, closing, or signature — the tool adds those.
    Example:
        I'm excited to apply for the Analytics Director role at Acme...

        Over the past nine years I've built and led analytics functions...

        I'd love to bring this experience to Acme's team...

Returns:
    Prints the Google Drive shareable URL to stdout.

Exit codes:
    0 = success
    1 = error (details on stderr)
"""

import argparse
import os
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.google_auth import get_credentials
from tools.drive_helpers import get_or_create_folder

TMP_FOLDER = Path(__file__).parent.parent / ".tmp"
ENV_PATH = Path(__file__).parent.parent / ".env"
COVER_LETTER_SUBFOLDER_NAME = "Cover Letters"


def _load_env() -> dict:
    env = {}
    if ENV_PATH.exists():
        for line in ENV_PATH.read_text().splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, _, v = line.partition("=")
                env[k.strip()] = v.strip().strip('"').strip("'")
    return env


def _safe_filename(text: str) -> str:
    return re.sub(r"[^\w\s-]", "", text).strip().replace(" ", "_")


def build_docx(company: str, job_title: str, paragraphs: list[str]) -> Path:
    """
    Create a clean cover letter .docx from the provided paragraphs.
    Returns the path to the saved file.
    """
    doc = Document()

    # Match resume margins exactly (0.75" all sides)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Default paragraph style — Arial to match resume
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Resume header colors
    _NAVY   = RGBColor(0x1B, 0x2A, 0x4A)
    _GRAY   = RGBColor(0x66, 0x66, 0x66)
    _BLUE   = RGBColor(0x2E, 0x6D, 0xA4)

    def add_para(text: str = "", bold: bool = False, size: float = 10,
                 color: RGBColor = None, align=None, space_after: int = 0) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color

    def add_contact_line(contact_str: str) -> None:
        """Reproduce resume contact-line color scheme: phone/separators gray, links blue."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = contact_str.split("  \u00b7  ")   # split on " · "
        for i, part in enumerate(parts):
            if i > 0:
                sep = p.add_run("  \u00b7  ")
                sep.font.name = "Arial"
                sep.font.size = Pt(9.5)
                sep.font.color.rgb = _GRAY
            r = p.add_run(part)
            r.font.name = "Arial"
            r.font.size = Pt(9.5)
            # Phone is gray; email and LinkedIn are blue
            r.font.color.rgb = _GRAY if i == 0 else _BLUE

    # Header: name — read from .env so each user's name appears correctly
    env = _load_env()
    user_name = env.get("USER_NAME", "Your Name")
    contact = env.get("USER_CONTACT", env.get("USER_EMAIL", "your@email.com"))

    add_para(user_name.upper(), bold=True, size=26, color=_NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_contact_line(contact)

    # Thin separator line (empty paragraph with bottom border look via spacing)
    add_para(space_after=6)

    # Date
    add_para(date.today().strftime("%B %-d, %Y"), size=10, color=_GRAY, space_after=10)

    # Addressee block — no em dashes
    add_para(company, size=10, space_after=2)
    add_para(f"Re: {job_title}", size=10, space_after=12)

    # Salutation
    add_para("Dear Hiring Team,", size=10, space_after=10)

    # Body paragraphs
    for para in paragraphs:
        para = para.strip()
        if para:
            add_para(para, size=10, space_after=10)

    # Closing
    add_para("Best regards,", size=10, space_after=18)
    add_para(env.get("USER_NAME", "Your Name"), size=10)

    # Save to .tmp/
    TMP_FOLDER.mkdir(exist_ok=True)
    filename = f"Watson_Dan_{_safe_filename(company)}_{_safe_filename(job_title)}_CL.docx"
    out_path = TMP_FOLDER / filename
    doc.save(str(out_path))
    return out_path


def upload_cover_letter(local_path: Path, drive_folder_id: str = None) -> str:
    """
    Upload the .docx to Drive in the Cover Letters subfolder.
    Returns a shareable URL.
    """
    creds = get_credentials()
    drive_svc = build("drive", "v3", credentials=creds)

    # Find or create the Cover Letters subfolder
    folder_id = get_or_create_folder(drive_svc, COVER_LETTER_SUBFOLDER_NAME, drive_folder_id)

    file_meta = {
        "name": local_path.name,
        "parents": [folder_id],
        "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    media = MediaFileUpload(
        str(local_path),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    uploaded = drive_svc.files().create(
        body=file_meta, media_body=media, fields="id"
    ).execute()

    file_id = uploaded["id"]

    # Make shareable
    drive_svc.permissions().create(
        fileId=file_id,
        body={"type": "anyone", "role": "reader"},
    ).execute()

    return f"https://drive.google.com/file/d/{file_id}/view"


def create_and_upload(company: str, job_title: str, content_file: str) -> str:
    """
    Full pipeline: read content file → build .docx → upload → return URL.
    """
    content = Path(content_file).read_text(encoding="utf-8")
    # Split into paragraphs on blank lines
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

    env = _load_env()
    drive_folder_id = env.get("DRIVE_FOLDER_ID") or None

    docx_path = build_docx(company, job_title, paragraphs)
    url = upload_cover_letter(docx_path, drive_folder_id)
    return url


def main():
    parser = argparse.ArgumentParser(description="Build and upload a cover letter.")
    parser.add_argument("--company", required=True)
    parser.add_argument("--job_title", required=True)
    parser.add_argument("--content_file", required=True)
    args = parser.parse_args()

    try:
        url = create_and_upload(args.company, args.job_title, args.content_file)
        print(url)
        sys.exit(0)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
